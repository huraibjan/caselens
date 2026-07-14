"""In-process document-processing pipeline.

This runs the full PDF → pages → chunks → embeddings → AI analysis pipeline
directly inside the API process as a background task, with no Temporal server
or separate worker required. It is the default (`DOCUMENT_PROCESSING_MODE=inline`)
and makes the whole app hostable as a single FastAPI service on a free tier.

Each step opens its own database session and is idempotent (safe to re-run),
so a document can be re-processed by calling /documents/{id}/analyze again.
"""

import uuid

import fitz  # PyMuPDF
import structlog
import tiktoken
from fastapi import BackgroundTasks
from sqlalchemy import select, update

from caselens.ai_gateway.providers import AllProvidersFailedError, get_analysis_llm_provider
from caselens.db.models import (
    Document,
    DocumentChunk,
    DocumentPage,
    DocumentStatus,
    Embedding,
    Matter,
)
from caselens.db.session import async_session_factory
from caselens.documents.analysis import merge_matter_metadata, run_case_analysis
from caselens.storage.backend import get_storage_backend

logger = structlog.get_logger()


def _extract_pdf_pages(data: bytes) -> list[str]:
    """One string per PDF page (PyMuPDF)."""
    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page in pdf:
            pages.append(page.get_text())
    return pages


def _extract_docx_pages(data: bytes) -> list[str]:
    """Word .docx text (paragraphs + table cells). Word has no fixed pages,
    so the whole document is returned as a single page. Lightweight: just
    unzips and reads XML — no rendering/OCR."""
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return ["\n".join(parts)]


def _extract_text_pages(data: bytes) -> list[str]:
    """Plain-text file as a single page."""
    return [data.decode("utf-8", errors="replace")]


async def _extract_text(document_id: uuid.UUID, storage_key: str) -> None:
    """Download the file and extract text per page into DocumentPage rows.

    Supports PDF (PyMuPDF), Word .docx (python-docx), and plain text — chosen by
    file signature with a filename-extension fallback.
    """
    storage = get_storage_backend()
    data = await storage.download(storage_key)
    key = storage_key.lower()

    if data.startswith(b"%PDF-") or key.endswith(".pdf"):
        page_texts = _extract_pdf_pages(data)
        method = "pymupdf"
    elif key.endswith(".docx") or data.startswith(b"PK\x03\x04"):
        # PK.. is the ZIP signature shared by all .docx files
        page_texts = _extract_docx_pages(data)
        method = "python-docx"
    elif key.endswith(".txt"):
        page_texts = _extract_text_pages(data)
        method = "plaintext"
    else:
        raise ValueError(f"Unsupported file type for extraction: {storage_key}")

    async with async_session_factory() as session:
        existing = await session.execute(
            select(DocumentPage.id).where(DocumentPage.document_id == document_id).limit(1)
        )
        if existing.scalar_one_or_none() is not None:
            return  # already extracted — idempotent

        for i, text_content in enumerate(page_texts):
            session.add(
                DocumentPage(
                    document_id=document_id,
                    page_number=i + 1,
                    text_content=text_content,
                    char_count=len(text_content),
                    extraction_method=method,
                    extraction_quality=1.0 if text_content.strip() else 0.0,
                )
            )

        await session.execute(
            update(Document)
            .where(Document.id == document_id)
            .values(page_count=len(page_texts), status=DocumentStatus.PROCESSING)
        )
        await session.commit()


async def _create_chunks(document_id: uuid.UUID) -> None:
    """Page-aware token chunking (512 tokens, 64 overlap) into DocumentChunk rows."""
    encoding = tiktoken.get_encoding("cl100k_base")
    target_size, overlap = 512, 64

    async with async_session_factory() as session:
        existing = await session.execute(
            select(DocumentChunk.id).where(DocumentChunk.document_id == document_id).limit(1)
        )
        if existing.scalar_one_or_none() is not None:
            return

        result = await session.execute(
            select(DocumentPage)
            .where(DocumentPage.document_id == document_id)
            .order_by(DocumentPage.page_number)
        )
        for page in result.scalars().all():
            text = page.text_content or ""
            tokens = encoding.encode(text)
            if not tokens:
                continue

            idx = chunk_index = 0
            while idx < len(tokens):
                chunk_tokens = tokens[idx : idx + target_size]
                chunk_text = encoding.decode(chunk_tokens)
                start_char = text.find(chunk_text[:30]) if len(chunk_text) > 30 else 0
                if start_char == -1:
                    start_char = 0
                session.add(
                    DocumentChunk(
                        document_id=document_id,
                        page_number=page.page_number,
                        chunk_index=chunk_index,
                        text_content=chunk_text,
                        token_count=len(chunk_tokens),
                        start_char=start_char,
                        end_char=start_char + len(chunk_text),
                        metadata_={},
                    )
                )
                chunk_index += 1
                if idx + target_size >= len(tokens):
                    break
                idx += target_size - overlap

        await session.commit()


async def _generate_embeddings(document_id: uuid.UUID) -> None:
    """Embed all chunks (batched) into Embedding rows."""
    from caselens.ai_gateway.providers import get_embedding_provider

    provider = get_embedding_provider()

    async with async_session_factory() as session:
        result = await session.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.page_number, DocumentChunk.chunk_index)
        )
        chunks = result.scalars().all()
        if not chunks:
            return

        chunk_ids = [c.id for c in chunks]
        existing = await session.execute(
            select(Embedding.id).where(Embedding.chunk_id.in_(chunk_ids)).limit(1)
        )
        if existing.scalar_one_or_none() is not None:
            return

        batch_size = 50
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            resp = await provider.embed([c.text_content for c in batch])
            for chunk, vector in zip(batch, resp.vectors, strict=False):
                session.add(
                    Embedding(
                        chunk_id=chunk.id,
                        vector=vector,
                        model_name=provider.model_name,
                        model_version="1.0.0",
                        dimensions=provider.dimensions,
                    )
                )
        await session.commit()


async def _analyze(document_id: uuid.UUID) -> None:
    """Generate the AI summary + structured case intelligence, mark READY."""
    llm = get_analysis_llm_provider()

    # Load the text in a short-lived session, run the (slow) LLM calls with no
    # DB connection held, then write results in a second session.
    async with async_session_factory() as session:
        result = await session.execute(select(Document.title).where(Document.id == document_id))
        title = result.scalar_one_or_none()
        if title is None:
            return

        chunk_result = await session.execute(
            select(DocumentChunk.text_content)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.page_number, DocumentChunk.id)
        )
        full_text = "\n\n".join(chunk_result.scalars().all())

    summary, metadata_dict = await run_case_analysis(llm, title, full_text)

    async with async_session_factory() as session:
        doc_result = await session.execute(select(Document).where(Document.id == document_id))
        doc = doc_result.scalar_one_or_none()
        if not doc:
            return

        doc.summary = summary
        doc.status = DocumentStatus.READY
        doc.metadata_ = metadata_dict

        matter_result = await session.execute(select(Matter).where(Matter.id == doc.matter_id))
        matter = matter_result.scalar_one_or_none()
        if matter:
            merged, new_title = merge_matter_metadata(matter.metadata_, metadata_dict)
            matter.metadata_ = merged
            if new_title:
                matter.title = new_title

        await session.commit()


async def _mark_error(document_id: uuid.UUID) -> None:
    async with async_session_factory() as session:
        await session.execute(
            update(Document).where(Document.id == document_id).values(status=DocumentStatus.ERROR)
        )
        await session.commit()


async def process_document(document_id_str: str) -> None:
    """Run the full pipeline for one document. Safe to call as a background task.

    On any failure the document is marked ERROR (retryable via re-analyze) and
    the error is logged — it never raises out of the background task.
    """
    document_id = uuid.UUID(document_id_str)
    logger.info("pipeline.start", document_id=document_id_str)
    try:
        await _extract_text(document_id, await _storage_key_for(document_id))
        await _create_chunks(document_id)
        await _generate_embeddings(document_id)
        await _analyze(document_id)
        logger.info("pipeline.done", document_id=document_id_str)
    except AllProvidersFailedError:
        logger.error("pipeline.ai_unavailable", document_id=document_id_str)
        await _mark_error(document_id)
    except Exception as e:
        logger.error("pipeline.failed", document_id=document_id_str, error=str(e))
        await _mark_error(document_id)


async def _storage_key_for(document_id: uuid.UUID) -> str:
    async with async_session_factory() as session:
        result = await session.execute(
            select(Document.storage_key).where(Document.id == document_id)
        )
        key = result.scalar_one_or_none()
        if not key:
            raise ValueError(f"Document {document_id} not found")
        return key


async def dispatch_processing(background_tasks: BackgroundTasks, doc: Document) -> None:
    """Kick off processing for a document using the configured mode.

    Default ("inline") schedules an in-process background task — no Temporal or
    separate worker required. "temporal" dispatches to the Temporal worker.
    """
    from caselens.config import settings

    if settings.DOCUMENT_PROCESSING_MODE == "temporal":
        await _dispatch_temporal(doc)
    else:
        background_tasks.add_task(process_document, str(doc.id))


async def _dispatch_temporal(doc: Document) -> None:
    from dataclasses import dataclass

    from temporalio.client import Client

    from caselens.config import settings

    @dataclass
    class DocumentProcessingInput:
        document_id: str
        organization_id: str
        matter_id: str
        storage_key: str
        original_filename: str

    client = await Client.connect(settings.TEMPORAL_HOST)
    await client.start_workflow(
        "DocumentProcessingWorkflow",
        DocumentProcessingInput(
            document_id=str(doc.id),
            organization_id=str(doc.organization_id),
            matter_id=str(doc.matter_id),
            storage_key=doc.storage_key,
            original_filename=doc.original_filename,
        ),
        id=f"doc-proc-{doc.id}-{uuid.uuid4().hex[:6]}",
        task_queue=settings.TEMPORAL_TASK_QUEUE,
    )
